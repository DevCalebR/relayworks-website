#!/usr/bin/env python3
"""Generate the public PDF and DOCX resumes from resume/resume.json."""

from __future__ import annotations

import json
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "resume" / "resume.json"
OUTPUT_DIR = ROOT / "public" / "resume"
PDF_PATH = OUTPUT_DIR / "Caleb_Rogers_Software_Engineer_Resume.pdf"
DOCX_PATH = OUTPUT_DIR / "Caleb_Rogers_Software_Engineer_Resume.docx"

INK = "#17212B"
MUTED = "#55616D"
ACCENT = "#126355"
LINE = "#D8DEE4"


def load_data() -> dict:
    return json.loads(SOURCE.read_text(encoding="utf-8"))


def add_docx_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "126355")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_border(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "D8DEE4")
    borders.append(bottom)
    properties.append(borders)


def configure_docx(document: Document, data: dict) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor(23, 33, 43)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.04

    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(23, 33, 43)
    styles["Title"].paragraph_format.space_after = Pt(1)

    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(18, 99, 85)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 1"].paragraph_format.space_before = Pt(7)
    styles["Heading 1"].paragraph_format.space_after = Pt(3)
    styles["Heading 2"].font.size = Pt(10.5)
    styles["Heading 2"].paragraph_format.space_before = Pt(4)
    styles["Heading 2"].paragraph_format.space_after = Pt(1)

    if "Resume Label" not in styles:
        label_style = styles.add_style("Resume Label", WD_STYLE_TYPE.PARAGRAPH)
    else:
        label_style = styles["Resume Label"]
    label_style.font.name = "Aptos"
    label_style.font.size = Pt(8)
    label_style.font.bold = True
    label_style.font.color.rgb = RGBColor(85, 97, 109)
    label_style.paragraph_format.space_after = Pt(1)

    props = document.core_properties
    props.title = f"{data['name']} - {data['title']} Resume"
    props.subject = "Software engineering resume"
    props.author = data["name"]
    props.keywords = "software engineer, Python, FastAPI, Next.js, React, TypeScript, APIs, webhooks"
    props.comments = ""


def add_docx_header(document: Document, data: dict) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(data["name"])
    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_after = Pt(2)
    run = role.add_run(data["title"])
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(18, 99, 85)

    location = document.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location.paragraph_format.space_after = Pt(2)
    location.add_run(data["location"])

    links = document.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(5)
    contact = data["contact"]
    for index, (label, url) in enumerate(
        (("Portfolio", contact["portfolio"]), ("GitHub", contact["github"]), ("LinkedIn", contact["linkedin"]))
    ):
        if index:
            links.add_run("  |  ")
        links.add_run(f"{label}: ")
        add_docx_hyperlink(links, url, url)
    set_cell_border(links)


def add_docx_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.first_line_indent = Inches(-0.13)
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run("- ").bold = True
    paragraph.add_run(text)


def add_docx_project(document: Document, project: dict) -> None:
    heading = document.add_paragraph(style="Heading 2")
    heading.add_run(project["name"])
    label = document.add_paragraph(style="Resume Label")
    label.add_run(f"{project['status']} | {project['stack']}")
    body = document.add_paragraph(project["description"], style="Normal")
    body.paragraph_format.keep_together = True
    if project["url"]:
        link = document.add_paragraph(style="Resume Label")
        link.add_run(f"{project['linkLabel']}: ")
        add_docx_hyperlink(link, project["url"], project["url"])


def generate_docx(data: dict) -> None:
    document = Document()
    configure_docx(document, data)
    add_docx_header(document, data)

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(data["summary"])

    document.add_heading("Technical Skills", level=1)
    for skill in data["skills"]:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.add_run(f"{skill['label']}: ").bold = True
        paragraph.add_run(skill["value"])

    document.add_heading("Professional Experience", level=1)
    for entry in data["experience"]:
        heading = document.add_paragraph(style="Heading 2")
        heading.add_run(f"{entry['organization']} | {entry['role']} | {entry['dates']}")
        for bullet in entry["bullets"]:
            add_docx_bullet(document, bullet)

    document.add_heading("Selected Engineering Work", level=1)
    for index, project in enumerate(data["projects"]):
        if index == 2:
            document.add_page_break()
            document.add_heading("Selected Engineering Work (continued)", level=1)
        add_docx_project(document, project)

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f"{data['name']} | {data['title']}")
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(85, 97, 109)

    document.save(DOCX_PATH)


def pdf_styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("Name", parent=base["Title"], fontName="Helvetica-Bold", fontSize=22, leading=24, alignment=TA_CENTER, textColor=colors.HexColor(INK), spaceAfter=1),
        "role": ParagraphStyle("Role", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=14, alignment=TA_CENTER, textColor=colors.HexColor(ACCENT), spaceAfter=2),
        "contact": ParagraphStyle("Contact", parent=base["Normal"], fontName="Helvetica", fontSize=8, leading=10.5, alignment=TA_CENTER, textColor=colors.HexColor(MUTED), spaceAfter=4, borderPadding=(0, 0, 5, 0), borderColor=colors.HexColor(LINE), borderWidth=0, borderBottomWidth=.6),
        "section": ParagraphStyle("Section", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=14, textColor=colors.HexColor(ACCENT), spaceBefore=6, spaceAfter=3, keepWithNext=True),
        "project": ParagraphStyle("Project", parent=base["Heading3"], fontName="Helvetica-Bold", fontSize=10.2, leading=12, textColor=colors.HexColor(INK), spaceBefore=4, spaceAfter=1, keepWithNext=True),
        "label": ParagraphStyle("Label", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=7.7, leading=9.3, textColor=colors.HexColor(MUTED), spaceAfter=1, keepWithNext=True),
        "body": ParagraphStyle("Body", parent=base["Normal"], fontName="Helvetica", fontSize=8.8, leading=11.1, textColor=colors.HexColor(INK), spaceAfter=2.5),
        "bullet": ParagraphStyle("Bullet", parent=base["Normal"], fontName="Helvetica", fontSize=8.8, leading=11.1, leftIndent=10, firstLineIndent=-8, textColor=colors.HexColor(INK), spaceAfter=2),
        "link": ParagraphStyle("Link", parent=base["Normal"], fontName="Helvetica", fontSize=7.7, leading=9.3, textColor=colors.HexColor(ACCENT), spaceAfter=2.5),
        "footer": ParagraphStyle("Footer", parent=base["Normal"], fontName="Helvetica", fontSize=7.5, textColor=colors.HexColor(MUTED), alignment=TA_CENTER),
    }


def pdf_link(label: str, url: str) -> str:
    safe_url = escape(url, {'"': '&quot;'})
    return f"{escape(label)}: <link href=\"{safe_url}\" color=\"{ACCENT}\"><u>{escape(url)}</u></link>"


def generate_pdf(data: dict) -> None:
    styles = pdf_styles()
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=0.57 * inch,
        leftMargin=0.57 * inch,
        topMargin=0.42 * inch,
        bottomMargin=0.42 * inch,
        title=f"{data['name']} - {data['title']} Resume",
        author=data["name"],
        subject="Software engineering resume",
        keywords="software engineer, Python, FastAPI, Next.js, React, TypeScript, APIs, webhooks",
    )
    story = [
        Paragraph(escape(data["name"]), styles["name"]),
        Paragraph(escape(data["title"]), styles["role"]),
        Paragraph(escape(data["location"]), styles["body"]),
    ]
    contact = data["contact"]
    contact_markup = " &nbsp; | &nbsp; ".join(
        pdf_link(label, url)
        for label, url in (("Portfolio", contact["portfolio"]), ("GitHub", contact["github"]), ("LinkedIn", contact["linkedin"]))
    )
    story.append(Paragraph(contact_markup, styles["contact"]))

    story.extend([Paragraph("Professional Summary", styles["section"]), Paragraph(escape(data["summary"]), styles["body"])])
    story.append(Paragraph("Technical Skills", styles["section"]))
    for skill in data["skills"]:
        story.append(Paragraph(f"<b>{escape(skill['label'])}:</b> {escape(skill['value'])}", styles["body"]))

    story.append(Paragraph("Professional Experience", styles["section"]))
    for entry in data["experience"]:
        story.append(Paragraph(escape(f"{entry['organization']} | {entry['role']} | {entry['dates']}"), styles["project"]))
        for bullet in entry["bullets"]:
            story.append(Paragraph(f"- {escape(bullet)}", styles["bullet"]))

    story.append(Paragraph("Selected Engineering Work", styles["section"]))
    for index, project in enumerate(data["projects"]):
        if index == 2:
            story.extend([PageBreak(), Paragraph("Selected Engineering Work (continued)", styles["section"])])
        story.extend(
            [
                Paragraph(escape(project["name"]), styles["project"]),
                Paragraph(escape(f"{project['status']} | {project['stack']}"), styles["label"]),
                Paragraph(escape(project["description"]), styles["body"]),
            ]
        )
        if project["url"]:
            story.append(Paragraph(pdf_link(project["linkLabel"], project["url"]), styles["link"]))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(MUTED))
        canvas.drawCentredString(A4[0] / 2, 0.22 * inch, f"{data['name']} | {data['title']} | Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    generate_docx(data)
    generate_pdf(data)
    print(PDF_PATH.relative_to(ROOT))
    print(DOCX_PATH.relative_to(ROOT))


if __name__ == "__main__":
    main()
